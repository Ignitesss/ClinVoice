# -*- coding: utf-8 -*-
"""Экспорт протокола в .docx / .txt (без Streamlit)."""

from __future__ import annotations

from datetime import datetime
from typing import Any, Dict
from zoneinfo import ZoneInfo

from docx import Document


def format_consultation_date_gmt3() -> str:
    return datetime.now(ZoneInfo("Europe/Moscow")).strftime("%d.%m.%Y %H:%M (GMT+3)")


def create_structured_protocol_docx(fields: Dict[str, Any], consultation_date: str) -> Document:
    doc = Document()
    doc.add_heading("Протокол консультации", 0)
    doc.add_paragraph(f"Дата: {consultation_date}")
    sections = [
        ("Жалобы", fields.get("complaints", "")),
        ("Анамнез", fields.get("anamnesis", "")),
        ("Заключение", fields.get("conclusion", "")),
        ("Рекомендации", fields.get("recommendations", "")),
    ]
    for title, body in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body if (body or "").strip() else "—")
    return doc


def build_structured_protocol_txt(fields: Dict[str, Any], consultation_date: str) -> str:
    parts = [
        "Протокол консультации",
        "",
        f"Дата: {consultation_date}",
        "",
        "Жалобы",
        fields.get("complaints", "") or "—",
        "",
        "Анамнез",
        fields.get("anamnesis", "") or "—",
        "",
        "Заключение",
        fields.get("conclusion", "") or "—",
        "",
        "Рекомендации",
        fields.get("recommendations", "") or "—",
        "",
    ]
    return "\n".join(parts)
