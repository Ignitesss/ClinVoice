# -*- coding: utf-8 -*-
"""
ClinVoice - Medical Speech Recognition App (врачебный сценарий: запись, транскрибация, протокол).
"""

import os

import streamlit as st
import streamlit.components.v1 as components

# Ключ `webrtc_streamer(..., key=...)` — должен совпадать с проверками в патче ниже.
CLINVOICE_WEBRTC_STREAMER_KEY = "clinvoice_webrtc_mic"


def _patch_streamlit_safe_session_state_for_webrtc_cache_clear() -> None:
    """
    После «Clear cache» / сброса session_state Streamlit вызывает on_change у streamlit-webrtc
    до выполнения скрипта; callback делает st.session_state[frontend_key] и падает с KeyError.
    Подставляем безопасные заглушки только для ключей нашего микрофона.
    """
    try:
        from streamlit.runtime.state import safe_session_state as sss
    except Exception:
        return

    cls = sss.SafeSessionState
    if getattr(cls, "_clinvoice_webrtc_ss_clear_patch", False):
        return

    def __getitem__(self, key: str):
        self._yield_callback()
        with self._lock:
            try:
                return self._state[key]
            except KeyError:
                if key == CLINVOICE_WEBRTC_STREAMER_KEY:
                    from streamlit_webrtc.component import (
                        WebRtcStreamerContext,
                        WebRtcStreamerState,
                    )

                    return WebRtcStreamerContext(
                        worker=None,
                        state=WebRtcStreamerState(playing=False, signalling=False),
                    )
                try:
                    from streamlit_webrtc.component import generate_frontend_component_key
                except Exception:
                    raise
                if key == generate_frontend_component_key(CLINVOICE_WEBRTC_STREAMER_KEY):
                    return {}
                raise

    cls.__getitem__ = __getitem__  # type: ignore[method-assign]
    cls._clinvoice_webrtc_ss_clear_patch = True


_patch_streamlit_safe_session_state_for_webrtc_cache_clear()


def resolve_app_cache_root() -> str:
    """Корень кэша артефактов: env CLINVOICE_CACHE_DIR, секрет Streamlit, или ~/.cache/clinvoice."""
    root = (os.environ.get("CLINVOICE_CACHE_DIR") or "").strip()
    if not root:
        try:
            if hasattr(st, "secrets") and st.secrets and "CLINVOICE_CACHE_DIR" in st.secrets:
                root = str(st.secrets["CLINVOICE_CACHE_DIR"]).strip()
        except Exception:
            pass
    if not root:
        root = os.path.join(os.path.expanduser("~"), ".cache", "clinvoice")
    os.makedirs(root, exist_ok=True)
    return root


def apply_disk_cache_layout(cache_root: str) -> None:
    """
    Свести загрузки к одному дереву (меньше разброса по ~/.cache).
    Hugging Face: HF_HOME / HF_HUB_CACHE; openai-whisper — отдельная подпапка (load_model).
    """
    hf_default = os.path.join(cache_root, "huggingface")
    openai_dir = os.path.join(cache_root, "openai-whisper")
    torch_dir = os.path.join(cache_root, "torch")
    for d in (hf_default, openai_dir, torch_dir):
        os.makedirs(d, exist_ok=True)
    os.environ.setdefault("HF_HOME", hf_default)
    hf_home = os.environ["HF_HOME"]
    os.makedirs(hf_home, exist_ok=True)
    hub = os.path.join(hf_home, "hub")
    os.environ.setdefault("HF_HUB_CACHE", hub)
    os.makedirs(os.environ["HF_HUB_CACHE"], exist_ok=True)
    os.environ.setdefault("TRANSFORMERS_CACHE", os.path.join(hf_home, "transformers"))
    os.makedirs(os.environ["TRANSFORMERS_CACHE"], exist_ok=True)
    os.environ.setdefault("TORCH_HOME", torch_dir)
    os.environ["_CLINVOICE_OPENAI_WHISPER_DIR"] = openai_dir


_APP_CACHE_ROOT = resolve_app_cache_root()
apply_disk_cache_layout(_APP_CACHE_ROOT)

import base64
import io
import json
import struct
import tempfile
import threading
import wave
import whisper
from docx import Document
import torch
from datetime import datetime, timedelta
from typing import List, Optional, Tuple
from zoneinfo import ZoneInfo

from streamlit_webrtc import RTCConfiguration, WebRtcMode, webrtc_streamer


def _patch_streamlit_webrtc_shutdown_observer_stop() -> None:
    """
    streamlit-webrtc: при вложенном вызове SessionShutdownObserver.stop() из polling-потока
    поле _polling_thread обнуляется до is_alive() в основном потоке → AttributeError.
    Сохраняем ссылку на Thread в локальной переменной (как следовало бы в апстриме).
    """
    import logging

    try:
        from streamlit_webrtc.shutdown import SessionShutdownObserver
    except Exception:
        return

    def _safe_stop(self, timeout: float = 1.0) -> None:
        poll_thread = self._polling_thread
        if not poll_thread:
            return
        self._polling_thread_stop_event.set()
        log = logging.getLogger("streamlit_webrtc.shutdown")
        if threading.current_thread() is not poll_thread:
            poll_thread.join(timeout=timeout)
            if poll_thread.is_alive():
                log.warning("ShutdownPolling thread did not exit cleanly")
            else:
                log.debug("ShutdownPolling thread stopped cleanly")
        else:
            log.debug("Stop called from polling thread itself, skipping join.")
        self._polling_thread = None

    SessionShutdownObserver.stop = _safe_stop  # type: ignore[method-assign]


_patch_streamlit_webrtc_shutdown_observer_stop()


def _patch_aioice_transaction_retry() -> None:
    """
    aioice: после закрытия UDP-транспорта таймер STUN всё ещё вызывает sendto → NoneType.
    Гасим повтор и завершаем future, чтобы не сыпались callback-ошибки в asyncio (особенно на Python 3.14).
    """
    import asyncio

    try:
        from aioice.stun import Transaction, TransactionTimeout
    except Exception:
        return
    if getattr(Transaction, "_clinvoice_retry_patched", False):
        return

    def _safe_retry(self) -> None:
        tries = getattr(self, "_Transaction__tries")
        tries_max = getattr(self, "_Transaction__tries_max")
        future = getattr(self, "_Transaction__future")
        if tries >= tries_max:
            if not future.done():
                try:
                    future.set_exception(TransactionTimeout())
                except (RuntimeError, asyncio.InvalidStateError):
                    pass
            return
        proto = getattr(self, "_Transaction__protocol")
        req = getattr(self, "_Transaction__request")
        addr = getattr(self, "_Transaction__addr")
        try:
            proto.send_stun(req, addr)
        except (AttributeError, OSError, RuntimeError):
            handle = getattr(self, "_Transaction__timeout_handle", None)
            if handle:
                try:
                    handle.cancel()
                except Exception:
                    pass
            setattr(self, "_Transaction__timeout_handle", None)
            if not future.done():
                try:
                    future.set_exception(TransactionTimeout())
                except (RuntimeError, asyncio.InvalidStateError):
                    pass
            return
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            if not future.done():
                try:
                    future.set_exception(TransactionTimeout())
                except (RuntimeError, asyncio.InvalidStateError):
                    pass
            return
        if loop.is_closed():
            if not future.done():
                try:
                    future.set_exception(TransactionTimeout())
                except (RuntimeError, asyncio.InvalidStateError):
                    pass
            return
        delay = getattr(self, "_Transaction__timeout_delay")
        setattr(
            self,
            "_Transaction__timeout_handle",
            loop.call_later(delay, _safe_retry, self),
        )
        setattr(self, "_Transaction__timeout_delay", delay * 2)
        setattr(self, "_Transaction__tries", tries + 1)

    Transaction._Transaction__retry = _safe_retry  # type: ignore[assignment]
    Transaction._clinvoice_retry_patched = True


_patch_aioice_transaction_retry()

from protocol import (
    PROTOCOL_FIELD_KEYS,
    fill_protocol_from_transcript,
    format_protocol_editor_text,
    parse_protocol_editor_text,
    resolve_yandex_api_key,
    resolve_yandex_folder_id,
    resolve_yandex_iam_token,
    yandex_llm_configured,
)
from webrtc_draft import DraftAudioProcessor, advance_after_empty_transcript

# For loading fine-tuned Whisper models
try:
    from transformers import AutoConfig, WhisperForConditionalGeneration, WhisperProcessor
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False

DEFAULT_HF_FINETUNED_REPO = "Ignites/fine_tuned_med_whisper_rus"
DEFAULT_ASR_CHUNK_SECONDS = 30.0


def resolve_live_whisper_interval_sec() -> float:
    """
    Период тика инкрементального live-Whisper (секунды).
    **CLINVOICE_LIVE_WHISPER_INTERVAL_SEC** (env или Streamlit Secrets).
    По умолчанию 6; значение ограничивается диапазоном [3, 15] (чаще тик — меньше «пропусков» в UI).
    """
    raw = (os.environ.get("CLINVOICE_LIVE_WHISPER_INTERVAL_SEC") or "").strip()
    if not raw:
        try:
            if hasattr(st, "secrets") and st.secrets and "CLINVOICE_LIVE_WHISPER_INTERVAL_SEC" in st.secrets:
                raw = str(st.secrets["CLINVOICE_LIVE_WHISPER_INTERVAL_SEC"]).strip()
        except Exception:
            pass
    default = 6.0
    if not raw:
        v = default
    else:
        try:
            v = float(raw)
        except ValueError:
            v = default
    return max(3.0, min(15.0, v))


def resolve_draft_tail_max_seconds() -> float:
    """
    Макс. длина одного PCM-среза для инкрементального live-черновика (сек).
    **CLINVOICE_DRAFT_TAIL_MAX_SECONDS** — env или Secrets; по умолчанию 15; диапазон [5, 60].
    """
    raw = (os.environ.get("CLINVOICE_DRAFT_TAIL_MAX_SECONDS") or "").strip()
    if not raw:
        try:
            if hasattr(st, "secrets") and st.secrets and "CLINVOICE_DRAFT_TAIL_MAX_SECONDS" in st.secrets:
                raw = str(st.secrets["CLINVOICE_DRAFT_TAIL_MAX_SECONDS"]).strip()
        except Exception:
            pass
    default = 15.0
    if not raw:
        v = default
    else:
        try:
            v = float(raw)
        except ValueError:
            v = default
    return max(5.0, min(60.0, v))


def resolve_draft_min_new_seconds() -> float:
    """
    Минимум нового PCM (сек) перед очередным куском live-черновика.
    **CLINVOICE_DRAFT_MIN_NEW_SECONDS**; по умолчанию 1.0; диапазон [0.2, 5].
    """
    raw = (os.environ.get("CLINVOICE_DRAFT_MIN_NEW_SECONDS") or "").strip()
    if not raw:
        try:
            if hasattr(st, "secrets") and st.secrets and "CLINVOICE_DRAFT_MIN_NEW_SECONDS" in st.secrets:
                raw = str(st.secrets["CLINVOICE_DRAFT_MIN_NEW_SECONDS"]).strip()
        except Exception:
            pass
    default = 1.0
    if not raw:
        v = default
    else:
        try:
            v = float(raw)
        except ValueError:
            v = default
    return max(0.2, min(5.0, v))


def resolve_draft_beam_size() -> int:
    """**CLINVOICE_DRAFT_BEAM_SIZE** для live-кусков (faster-whisper). По умолчанию 1; [1, 5]."""
    raw = (os.environ.get("CLINVOICE_DRAFT_BEAM_SIZE") or "").strip()
    if not raw:
        try:
            if hasattr(st, "secrets") and st.secrets and "CLINVOICE_DRAFT_BEAM_SIZE" in st.secrets:
                raw = str(st.secrets["CLINVOICE_DRAFT_BEAM_SIZE"]).strip()
        except Exception:
            pass
    if raw.isdigit():
        return max(1, min(5, int(raw)))
    return 1


def resolve_draft_vad_filter() -> bool:
    """
    **CLINVOICE_DRAFT_VAD_FILTER**: 1/true/yes — включить VAD для live-кусков (faster-whisper).
    Иначе выключено. По умолчанию включено.
    """
    raw = (os.environ.get("CLINVOICE_DRAFT_VAD_FILTER") or "").strip().lower()
    if not raw:
        try:
            if hasattr(st, "secrets") and st.secrets and "CLINVOICE_DRAFT_VAD_FILTER" in st.secrets:
                raw = str(st.secrets["CLINVOICE_DRAFT_VAD_FILTER"]).strip().lower()
        except Exception:
            pass
    if raw in ("0", "false", "no", "off"):
        return False
    if raw in ("1", "true", "yes", "on"):
        return True
    return True


def resolve_draft_tail_overlap_sec() -> float:
    """
    Перекрытие соседних кусков PCM для контекста (сек). **CLINVOICE_DRAFT_TAIL_OVERLAP_SEC**; [0, 3].
    По умолчанию 0.35 с (стыки сегментов без перекрытия часто «съедают» слова).
    """
    raw = (os.environ.get("CLINVOICE_DRAFT_TAIL_OVERLAP_SEC") or "").strip()
    if not raw:
        try:
            if hasattr(st, "secrets") and st.secrets and "CLINVOICE_DRAFT_TAIL_OVERLAP_SEC" in st.secrets:
                raw = str(st.secrets["CLINVOICE_DRAFT_TAIL_OVERLAP_SEC"]).strip()
        except Exception:
            pass
    if not raw:
        return 0.35
    try:
        v = float(raw)
    except ValueError:
        return 0.35
    return max(0.0, min(3.0, v))


def _secret_str(name: str) -> str:
    try:
        if hasattr(st, "secrets") and st.secrets and name in st.secrets:
            return str(st.secrets[name]).strip()
    except Exception:
        pass
    return ""


def _resolve_faster_whisper_no_speech_threshold(*, draft: bool) -> float:
    """
    Порог «нет речи» для faster-whisper (выше — меньше галлюцинаций на тишине).
    **CLINVOICE_DRAFT_NO_SPEECH_THRESHOLD** / **CLINVOICE_FINAL_NO_SPEECH_THRESHOLD** или общий
    **CLINVOICE_WHISPER_NO_SPEECH_THRESHOLD** (env / Secrets). Диапазон [0.35, 0.95].
    """
    raw = (os.environ.get("CLINVOICE_WHISPER_NO_SPEECH_THRESHOLD") or "").strip()
    if not raw:
        raw = _secret_str("CLINVOICE_WHISPER_NO_SPEECH_THRESHOLD")
    if draft:
        dr = (os.environ.get("CLINVOICE_DRAFT_NO_SPEECH_THRESHOLD") or "").strip() or _secret_str(
            "CLINVOICE_DRAFT_NO_SPEECH_THRESHOLD"
        )
        if dr:
            raw = dr
    else:
        fn = (os.environ.get("CLINVOICE_FINAL_NO_SPEECH_THRESHOLD") or "").strip() or _secret_str(
            "CLINVOICE_FINAL_NO_SPEECH_THRESHOLD"
        )
        if fn:
            raw = fn
    default = 0.66 if draft else 0.62
    if not raw:
        v = default
    else:
        try:
            v = float(raw)
        except ValueError:
            v = default
    return max(0.35, min(0.95, v))


def _resolve_faster_whisper_compression_ratio(*, draft: bool) -> float:
    """
    Порог compression ratio (ниже — жёстче отсекать зацикленные галлюцинации).
    **CLINVOICE_WHISPER_COMPRESSION_RATIO_THRESHOLD** или draft/final-суффиксы; [1.2, 3.5].
    """
    raw = (os.environ.get("CLINVOICE_WHISPER_COMPRESSION_RATIO_THRESHOLD") or "").strip()
    if not raw:
        raw = _secret_str("CLINVOICE_WHISPER_COMPRESSION_RATIO_THRESHOLD")
    if draft:
        dr = (os.environ.get("CLINVOICE_DRAFT_COMPRESSION_RATIO_THRESHOLD") or "").strip() or _secret_str(
            "CLINVOICE_DRAFT_COMPRESSION_RATIO_THRESHOLD"
        )
        if dr:
            raw = dr
    else:
        fn = (os.environ.get("CLINVOICE_FINAL_COMPRESSION_RATIO_THRESHOLD") or "").strip() or _secret_str(
            "CLINVOICE_FINAL_COMPRESSION_RATIO_THRESHOLD"
        )
        if fn:
            raw = fn
    default = 2.0 if draft else 2.35
    if not raw:
        v = default
    else:
        try:
            v = float(raw)
        except ValueError:
            v = default
    return max(1.2, min(3.5, v))


def _resolve_whisper_initial_prompt() -> str:
    """Опциональный **CLINVOICE_WHISPER_INITIAL_PROMPT** — смещение к домену (короткая строка)."""
    raw = (os.environ.get("CLINVOICE_WHISPER_INITIAL_PROMPT") or "").strip()
    if not raw:
        raw = _secret_str("CLINVOICE_WHISPER_INITIAL_PROMPT")
    return (raw[:224] if raw else "").strip()


def _resolve_draft_min_pcm_rms() -> float:
    """
    Если > 0, live-куски с RMS ниже порога не отправляются в Whisper (тишина).
    **CLINVOICE_DRAFT_MIN_PCM_RMS**; 0 = выключено. Типичный старт 80–200 для int16.
    """
    raw = (os.environ.get("CLINVOICE_DRAFT_MIN_PCM_RMS") or "").strip()
    if not raw:
        raw = _secret_str("CLINVOICE_DRAFT_MIN_PCM_RMS")
    if not raw:
        return 0.0
    try:
        v = float(raw)
    except ValueError:
        return 0.0
    return max(0.0, min(5000.0, v))


def _pcm_s16le_mono_rms(pcm: bytes) -> float:
    """RMS по int16 little-endian mono."""
    if len(pcm) < 2:
        return 0.0
    n = len(pcm) // 2
    if n <= 0:
        return 0.0
    sum_sq = 0.0
    for i in range(0, n * 2, 2):
        s = struct.unpack_from("<h", pcm, i)[0]
        sum_sq += float(s) * float(s)
    return (sum_sq / float(n)) ** 0.5


def _ice_server_key(entry: dict) -> str:
    u = entry.get("urls")
    if isinstance(u, list):
        return "|".join(sorted(str(x) for x in u))
    return str(u)


def resolve_webrtc_rtc_configuration() -> RTCConfiguration:
    """
    ICE для браузера и для aiortc на сервере: сначала встроенный набор streamlit-webrtc
    (Twilio / Hugging Face TURN при заданных TWILIO_* или HF_TOKEN), затем дополнительные STUN,
    затем JSON из CLINVOICE_WEBRTC_ICE_SERVERS_JSON (env или Streamlit Secrets).
    """
    from streamlit_webrtc.credentials import get_available_ice_servers

    merged: List[dict] = []
    seen: set[str] = set()

    def add(entry: Optional[dict]) -> None:
        if not entry or "urls" not in entry:
            return
        k = _ice_server_key(entry)
        if k in seen:
            return
        seen.add(k)
        merged.append(dict(entry))

    try:
        for s in get_available_ice_servers():
            if isinstance(s, dict):
                add(s)
            else:
                add(dict(s))
    except Exception:
        pass

    for u in (
        "stun:stun.l.google.com:19302",
        "stun:stun1.l.google.com:19302",
        "stun:stun2.l.google.com:19302",
        "stun:stun3.l.google.com:19302",
        "stun:stun4.l.google.com:19302",
        "stun:stun.cloudflare.com:3478",
        "stun:stun.services.mozilla.com:3478",
        "stun:global.stun.twilio.com:3478",
    ):
        add({"urls": u})

    raw = (os.environ.get("CLINVOICE_WEBRTC_ICE_SERVERS_JSON") or "").strip()
    if not raw:
        try:
            if hasattr(st, "secrets") and st.secrets and "CLINVOICE_WEBRTC_ICE_SERVERS_JSON" in st.secrets:
                raw = str(st.secrets["CLINVOICE_WEBRTC_ICE_SERVERS_JSON"]).strip()
        except Exception:
            pass
    if raw:
        try:
            extra = json.loads(raw)
            if isinstance(extra, list):
                for item in extra:
                    if isinstance(item, dict):
                        add(item)
        except json.JSONDecodeError:
            pass

    if not merged:
        merged = [{"urls": "stun:stun.l.google.com:19302"}]

    return RTCConfiguration({"iceServers": merged})


WEBRTC_UI_RU: dict = {
    "start": "Начать запись",
    "stop": "Остановить запись",
    "select_device": "Выбрать микрофон",
    "media_api_not_available": "Медиа API недоступно в этом браузере",
    "device_ask_permission": "Разрешите доступ к микрофону",
    "device_not_available": "Микрофон недоступен",
    "device_access_denied": "Доступ к микрофону запрещён",
}


def ensure_webrtc_shared_initialized() -> None:
    """
    Восстанавливает webrtc_shared после очистки session state или при перезапуске
    только @st.fragment(run_every=...), когда верх скрипта не выполняется.
    """
    w = st.session_state.get("webrtc_shared")
    if not isinstance(w, dict):
        st.session_state.webrtc_shared = {
            "lock": threading.Lock(),
            "asr_lock": threading.Lock(),
            "pcm_accum": bytearray(),
            "live_whisper_text": "",
            "live_whisper_error": None,
            "live_draft_pcm_committed": 0,
            "_draft_empty_streak": 0,
        }
        return
    w.setdefault("lock", threading.Lock())
    w.setdefault("pcm_accum", bytearray())
    w.setdefault("asr_lock", threading.Lock())
    w.setdefault("live_whisper_text", "")
    w.setdefault("live_whisper_error", None)
    w.setdefault("live_draft_pcm_committed", 0)
    w.setdefault("_draft_empty_streak", 0)
    w.pop("live_whisper_last_processed_pcm_len", None)


def get_cached_asr_transcriber(model_size: str, hub_model_id: str) -> "AudioTranscriberWithMetrics":
    """Один экземпляр распознавателя на (hub, model_size) для live-кусков и опционального полного прогона."""
    safe = (hub_model_id or "").replace("/", "_")
    cache_key = f"_clinvoice_asr_{safe}_{model_size}"
    if cache_key not in st.session_state:
        st.session_state[cache_key] = AudioTranscriberWithMetrics(
            model_size=model_size, hub_model_id=hub_model_id, silent_ui=True
        )
    return st.session_state[cache_key]


def transcribe_pcm_bytes_to_text(
    transcriber: "AudioTranscriberWithMetrics",
    pcm: bytes,
    *,
    draft: bool,
) -> str:
    """PCM s16le mono 16k → временный WAV → Whisper (короткие срезы обычно без чанкования)."""
    if not pcm:
        return ""
    if draft:
        rms_floor = _resolve_draft_min_pcm_rms()
        if rms_floor > 0.0 and _pcm_s16le_mono_rms(pcm) < rms_floor:
            return ""
    merged_path = None
    try:
        fd, merged_path = tempfile.mkstemp(suffix=".wav")
        os.close(fd)
        with open(merged_path, "wb") as wf:
            wf.write(pcm_mono_s16le_to_wav_bytes(pcm))
        return transcribe_wav_in_chunks(transcriber, merged_path, language="ru", draft=draft)
    finally:
        if merged_path and os.path.isfile(merged_path):
            try:
                os.remove(merged_path)
            except OSError:
                pass


def flush_incremental_pcm_tail(
    shared: dict,
    transcriber: "AudioTranscriberWithMetrics",
    asr_lock: threading.Lock,
    max_segment_bytes: int,
    overlap_bytes: int,
) -> Optional[str]:
    """
    Догоняет необработанный хвост PCM в live_whisper_text (те же сегменты, что в фоне).
    Возвращает сообщение об ошибке или None.
    """
    lk = shared.get("lock")
    if not lk:
        return "Нет lock"
    min_new_b = int(resolve_draft_min_new_seconds() * 32000)
    while True:
        with lk:
            pcm = bytes(shared.get("pcm_accum") or b"")
            committed = int(shared.get("live_draft_pcm_committed") or 0)
            n = len(pcm)
        if committed >= n:
            break
        take = min(n - committed, max_segment_bytes)
        ov = min(overlap_bytes, committed)
        start = committed - ov
        chunk = pcm[start : committed + take]
        try:
            with asr_lock:
                text = transcribe_pcm_bytes_to_text(transcriber, chunk, draft=True)
        except Exception as e:
            with lk:
                shared["live_whisper_error"] = str(e)
            return str(e)
        t = (text or "").strip()
        with lk:
            shared["live_whisper_error"] = None
            prev = (shared.get("live_whisper_text") or "").strip()
            if t:
                shared["_draft_empty_streak"] = 0
                if prev:
                    shared["live_whisper_text"] = (prev + " " + t).strip()
                else:
                    shared["live_whisper_text"] = t
                shared["live_draft_pcm_committed"] = committed + take
            else:
                es = int(shared.get("_draft_empty_streak") or 0)
                adv, nes = advance_after_empty_transcript(chunk, take, min_new_b, es)
                if adv:
                    shared["_draft_empty_streak"] = 0
                    shared["live_draft_pcm_committed"] = committed + take
                else:
                    shared["_draft_empty_streak"] = nes
    return None


def build_webrtc_processor_factory(
    transcriber: "AudioTranscriberWithMetrics",
    interval_sec: float,
):
    """
    Вызывать из основного потока Streamlit. Фабрика без обращения к session_state
    внутри worker WebRTC — только замыкание на готовый shared и transcriber.
    """
    ensure_webrtc_shared_initialized()
    shared = st.session_state.webrtc_shared
    asr_lock = shared["asr_lock"]
    max_seg_b = int(resolve_draft_tail_max_seconds() * 32000)
    min_new_b = int(resolve_draft_min_new_seconds() * 32000)
    overlap_b = int(resolve_draft_tail_overlap_sec() * 32000)

    def _transcribe_pcm(pcm_slice: bytes) -> Tuple[str, Optional[str]]:
        if not pcm_slice:
            return "", None
        try:
            with asr_lock:
                text = transcribe_pcm_bytes_to_text(transcriber, pcm_slice, draft=True)
            return text, None
        except Exception as e:
            return "", str(e)

    def _factory():
        return DraftAudioProcessor(
            shared,
            _transcribe_pcm,
            interval_sec,
            max_seg_b,
            min_new_b,
            overlap_b,
        )

    return _factory


def resolve_asr_chunk_seconds() -> float:
    raw = (os.environ.get("CLINVOICE_ASR_CHUNK_SECONDS") or "").strip()
    if raw:
        try:
            v = float(raw)
            if v > 0:
                return v
        except ValueError:
            pass
    try:
        if hasattr(st, "secrets") and st.secrets and "CLINVOICE_ASR_CHUNK_SECONDS" in st.secrets:
            v = float(str(st.secrets["CLINVOICE_ASR_CHUNK_SECONDS"]).strip())
            if v > 0:
                return v
    except Exception:
        pass
    return DEFAULT_ASR_CHUNK_SECONDS


def pcm_mono_s16le_to_wav_bytes(pcm: bytes, sample_rate: int = 16000) -> bytes:
    """Обёртка сырых PCM s16le mono в WAV (для Whisper)."""
    out = io.BytesIO()
    with wave.open(out, "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(sample_rate)
        w.writeframes(pcm)
    return out.getvalue()


def transcribe_wav_in_chunks(
    transcriber: "AudioTranscriberWithMetrics",
    wav_path: str,
    language: str = "ru",
    *,
    draft: bool = False,
) -> str:
    """Длинный WAV режется на части ≤ resolve_asr_chunk_seconds(); короткий — один вызов transcribe_audio."""
    chunk_sec = resolve_asr_chunk_seconds()
    with wave.open(wav_path, "rb") as w:
        ch = w.getnchannels()
        sw = w.getsampwidth()
        fr = w.getframerate()
        nframes = w.getnframes()
        if ch != 1:
            raise ValueError("Ожидается моно WAV")
        frames_bytes = w.readframes(nframes)
    duration = nframes / float(fr)
    if duration <= chunk_sec:
        return transcriber.transcribe_audio(wav_path, language=language, draft=draft)

    chunk_frames = max(1, int(fr * chunk_sec))
    parts: List[str] = []
    offset = 0
    while offset < nframes:
        take = min(chunk_frames, nframes - offset)
        start_byte = offset * ch * sw
        end_byte = (offset + take) * ch * sw
        chunk_data = frames_bytes[start_byte:end_byte]
        tmp_path = None
        try:
            fd, tmp_path = tempfile.mkstemp(suffix=".wav")
            os.close(fd)
            with wave.open(tmp_path, "wb") as cw:
                cw.setnchannels(ch)
                cw.setsampwidth(sw)
                cw.setframerate(fr)
                cw.writeframes(chunk_data)
            parts.append(transcriber.transcribe_audio(tmp_path, language=language, draft=draft))
        finally:
            if tmp_path and os.path.isfile(tmp_path):
                try:
                    os.remove(tmp_path)
                except OSError:
                    pass
        offset += take
    return " ".join(p.strip() for p in parts if p.strip()).strip()


def resolve_hub_model_id() -> str:
    env_id = (os.environ.get("CLINVOICE_HF_MODEL_REPO") or "").strip()
    if env_id:
        return env_id
    try:
        if hasattr(st, "secrets") and "CLINVOICE_HF_MODEL_REPO" in st.secrets:
            return str(st.secrets["CLINVOICE_HF_MODEL_REPO"]).strip()
    except Exception:
        pass
    return DEFAULT_HF_FINETUNED_REPO


def resolve_whisper_engine() -> str:
    """
    Репозиторий на HF: «faster_whisper» (CTranslate2: model.bin) или «transformers»
    (PyTorch: model.safetensors / pytorch_model.bin). Задаётся CLINVOICE_WHISPER_ENGINE
    в env или Streamlit Secrets. По умолчанию — faster_whisper (репозиторий по умолчанию
    в формате CT2). Для старых репо только с PyTorch укажите: transformers | pytorch | hf.
    """
    raw = (os.environ.get("CLINVOICE_WHISPER_ENGINE") or "").strip().lower()
    if raw in ("transformers", "pytorch", "hf", "huggingface"):
        return "transformers"
    if raw in ("faster_whisper", "faster-whisper", "ct2", "ctranslate2"):
        return "faster_whisper"
    try:
        if hasattr(st, "secrets") and "CLINVOICE_WHISPER_ENGINE" in st.secrets:
            v = str(st.secrets["CLINVOICE_WHISPER_ENGINE"]).strip().lower()
            if v in ("transformers", "pytorch", "hf", "huggingface"):
                return "transformers"
            if v in ("faster_whisper", "faster-whisper", "ct2", "ctranslate2"):
                return "faster_whisper"
    except Exception:
        pass
    return "faster_whisper"


def infer_whisper_processor_repo(hub_model_id: str) -> str:
    """Базовый openai/whisper-* для процессора, если в finetune-репо нет preprocessor_config.json / tokenizer."""
    explicit = (os.environ.get("CLINVOICE_WHISPER_BASE_REPO") or "").strip()
    if explicit:
        return explicit
    try:
        cfg = AutoConfig.from_pretrained(hub_model_id)
        cand = getattr(cfg, "_name_or_path", None) or getattr(cfg, "name_or_path", None)
        if cand:
            s = str(cand).strip()
            if s.startswith("openai/whisper-"):
                return s
    except Exception:
        pass
    return "openai/whisper-small"


def openai_whisper_download_dir() -> str:
    return os.environ.get("_CLINVOICE_OPENAI_WHISPER_DIR") or os.path.join(
        _APP_CACHE_ROOT, "openai-whisper"
    )


def hf_hub_download_dir() -> str:
    return os.environ.get("HF_HUB_CACHE") or os.path.join(
        os.environ.get("HF_HOME", os.path.join(_APP_CACHE_ROOT, "huggingface")),
        "hub",
    )


# Не через @st.cache_resource: пункт меню «Clear cache» сбрасывает только st.cache_*,
# а экземпляр распознавателя живёт в session_state — после очистки оставалась «живая»
# ссылка на уже уничтоженную модель и запись переставала накапливаться.
_clinvoice_model_load_lock = threading.Lock()
_clinvoice_fw_whisper_models: dict[tuple[str, str, str], object] = {}
_clinvoice_openai_whisper_models: dict[tuple[str, str], object] = {}


def _load_faster_whisper_cached(repo_id: str, device_kw: str, compute_type: str):
    from faster_whisper import WhisperModel

    k = (str(repo_id), str(device_kw), str(compute_type))
    with _clinvoice_model_load_lock:
        if k not in _clinvoice_fw_whisper_models:
            _clinvoice_fw_whisper_models[k] = WhisperModel(
                repo_id,
                device=device_kw,
                compute_type=compute_type,
                download_root=hf_hub_download_dir(),
            )
        return _clinvoice_fw_whisper_models[k]


def _load_openai_whisper_cached(model_size: str):
    dev = "cuda" if torch.cuda.is_available() else "cpu"
    k = (str(model_size), dev)
    with _clinvoice_model_load_lock:
        if k not in _clinvoice_openai_whisper_models:
            _clinvoice_openai_whisper_models[k] = whisper.load_model(
                model_size, device=dev, download_root=openai_whisper_download_dir()
            )
        return _clinvoice_openai_whisper_models[k]


# Page config
st.set_page_config(
    page_title="ClinVoice",
    page_icon="🏥",
    layout="wide"
)

# До любых виджетов: WebRTC вызывает audio_processor_factory из фонового потока,
# там нельзя обращаться к st.session_state — только замыкание на готовый dict.
ensure_webrtc_shared_initialized()

if "live_transcript_editor" not in st.session_state:
    st.session_state.live_transcript_editor = ""
if "live_transcript_pause_auto_sync" not in st.session_state:
    st.session_state.live_transcript_pause_auto_sync = False

# ============ CLASSES FROM YOUR COLAB ============


class AudioTranscriberWithMetrics:
    def __init__(
        self,
        model_size="base",
        hub_model_id: Optional[str] = None,
        *,
        silent_ui: bool = False,
    ):
        """Whisper: openai-whisper, HF (PyTorch/transformers) или HF (CTranslate2 / faster-whisper)."""
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        engine = resolve_whisper_engine() if hub_model_id else "openai"
        safe_hub = (hub_model_id or "").replace("/", "_")
        transformers_cache_key = f"asr_transformers_{model_size}_{safe_hub or 'openai_whisper'}"
        self.use_faster_whisper = False
        self.use_transformers = False

        if hub_model_id and engine == "faster_whisper":
            if not silent_ui:
                st.info(
                    f"Загрузка модели (формат CTranslate2) с Hugging Face: {hub_model_id} "
                    f"(первый запуск может занять несколько минут)..."
                )
            device_kw = "cuda" if self.device.type == "cuda" else "cpu"
            compute_type = "float16" if device_kw == "cuda" else "int8"
            try:
                self.faster_model = _load_faster_whisper_cached(
                    hub_model_id, device_kw, compute_type
                )
            except ImportError:
                st.error("Нужен пакет faster-whisper. Выполните: pip install -r requirements.txt")
                st.stop()
            except Exception as e:
                st.error(f"Ошибка загрузки модели с Hub: {e}")
                st.stop()
            self.use_faster_whisper = True
            self.use_transformers = False
            self._draft_beam_size = resolve_draft_beam_size()
            self._draft_vad_filter = resolve_draft_vad_filter()
            self._fw_ns_draft = _resolve_faster_whisper_no_speech_threshold(draft=True)
            self._fw_ns_final = _resolve_faster_whisper_no_speech_threshold(draft=False)
            self._fw_cr_draft = _resolve_faster_whisper_compression_ratio(draft=True)
            self._fw_cr_final = _resolve_faster_whisper_compression_ratio(draft=False)
            self._fw_initial_prompt = _resolve_whisper_initial_prompt()
        elif hub_model_id:
            if not TRANSFORMERS_AVAILABLE:
                st.error("Пакет transformers необходим для дообученной модели с Hub (режим PyTorch).")
                st.stop()
            if transformers_cache_key not in st.session_state:
                if not silent_ui:
                    st.info(
                        f"Загрузка дообученной модели (PyTorch) с Hugging Face: {hub_model_id} "
                        f"(первый запуск может занять несколько минут)..."
                    )
                try:
                    model = WhisperForConditionalGeneration.from_pretrained(
                        hub_model_id,
                        torch_dtype=torch.float32,
                    )
                    try:
                        processor = WhisperProcessor.from_pretrained(hub_model_id)
                    except Exception:
                        base_proc = infer_whisper_processor_repo(hub_model_id)
                        if not silent_ui:
                            st.warning(
                                f"В репозитории «{hub_model_id}» нет полного процессора "
                                f"(нужен preprocessor_config.json и файлы токенизатора). "
                                f"Загружаю процессор с **{base_proc}**. "
                                f"Если размер не совпадает с дообучением, задайте переменную **CLINVOICE_WHISPER_BASE_REPO** "
                                f"(например, `openai/whisper-base`)."
                            )
                        processor = WhisperProcessor.from_pretrained(base_proc)
                    st.session_state[transformers_cache_key] = {
                        "model": model,
                        "processor": processor,
                        "feature_extractor": processor.feature_extractor,
                        "tokenizer": processor.tokenizer,
                    }
                except Exception as e:
                    st.error(f"Ошибка загрузки модели с Hub: {e}")
                    st.stop()
            cached = st.session_state[transformers_cache_key]
            self.model = cached["model"]
            self.processor = cached["processor"]
            self.feature_extractor = cached["feature_extractor"]
            self.tokenizer = cached["tokenizer"]
            self.use_transformers = True
            self.use_faster_whisper = False
            self.model.to(self.device)
            self.model.eval()
            self._draft_beam_size = resolve_draft_beam_size()
            self._draft_vad_filter = resolve_draft_vad_filter()
            self._fw_ns_draft = 0.6
            self._fw_ns_final = 0.6
            self._fw_cr_draft = 2.4
            self._fw_cr_final = 2.4
            self._fw_initial_prompt = _resolve_whisper_initial_prompt()
        else:
            if not silent_ui:
                st.info(f"Загрузка базовой модели Whisper ({model_size})...")
            self.model = _load_openai_whisper_cached(model_size)
            self.use_transformers = False
            self.use_faster_whisper = False
            self._draft_beam_size = 1
            self._draft_vad_filter = False
            self._fw_ns_draft = 0.6
            self._fw_ns_final = 0.6
            self._fw_cr_draft = 2.4
            self._fw_cr_final = 2.4
            self._fw_initial_prompt = ""

    def transcribe_audio(self, audio_path, language="ru", *, draft: bool = False):
        """Транскрибация аудиофайла. draft=True — быстрые настройки для live-кусков (только faster-whisper)."""
        if getattr(self, "use_faster_whisper", False):
            beam = getattr(self, "_draft_beam_size", 1) if draft else 5
            vad = getattr(self, "_draft_vad_filter", True) if draft else False
            ns = getattr(self, "_fw_ns_draft" if draft else "_fw_ns_final", 0.62)
            cr = getattr(self, "_fw_cr_draft" if draft else "_fw_cr_final", 2.35)
            ip = (getattr(self, "_fw_initial_prompt", None) or "").strip()
            kw = dict(
                language=language,
                beam_size=beam,
                vad_filter=vad,
                no_speech_threshold=ns,
                compression_ratio_threshold=cr,
                condition_on_previous_text=False if draft else True,
            )
            if ip:
                kw["initial_prompt"] = ip
            segments, _info = self.faster_model.transcribe(audio_path, **kw)
            return "".join(seg.text for seg in segments).strip()

        if self.use_transformers:
            # Use Transformers for fine-tuned model
            import librosa
            audio, sr = librosa.load(audio_path, sr=16000)
            dev = next(self.model.parameters()).device
            inputs = self.feature_extractor(
                audio, sampling_rate=16000, return_tensors="pt"
            )
            input_features = inputs["input_features"].to(dev)
            forced_decoder_ids = self.processor.get_decoder_prompt_ids(language=language, task="transcribe")
            with torch.no_grad():
                predicted_ids = self.model.generate(
                    input_features, forced_decoder_ids=forced_decoder_ids
                )
            transcription = self.processor.batch_decode(predicted_ids, skip_special_tokens=True)[0]
            return transcription
        else:
            # Use openai-whisper
            result = self.model.transcribe(audio_path, language=language)
            return result["text"]


# ============ DOCX / TXT GENERATION ============

def format_consultation_date_gmt3() -> str:
    """Дата/время в поясе Europe/Moscow (UTC+3)."""
    return datetime.now(ZoneInfo("Europe/Moscow")).strftime("%d.%m.%Y %H:%M (GMT+3)")


def create_structured_protocol_docx(fields: dict, consultation_date: str):
    """Протокол: дата → жалобы → анамнез → заключение → рекомендации."""
    doc = Document()
    doc.add_heading("Протокол консультации", 0)
    doc.add_paragraph(f"Дата: {consultation_date}")
    sections = [
        ("Жалобы", fields.get("complaints", "")),
        ("Анамнез", fields.get("anamnesis", "")),
        ("Заключение", fields.get("conclusion", "")),
        ("Рекомендации", fields.get("recommendations", "")),
    ]
    for title, body in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body if (body or "").strip() else "—")
    return doc


def trigger_browser_text_download(filename: str, text: str) -> None:
    """
    Пытается скачать .txt из контекста родительского фрейма Streamlit (в iframe часто блокируют клик).
    """
    b64 = base64.b64encode(text.encode("utf-8")).decode("ascii")
    fname_js = json.dumps(filename)
    b64_js = json.dumps(b64)
    components.html(
        f"""
        <script>
            (function () {{
                const fname = {fname_js};
                const b64 = {b64_js};
                const dataUrl = "data:text/plain;charset=utf-8;base64," + b64;
                function tryDownload(doc) {{
                    try {{
                        if (!doc || !doc.body) return false;
                        const a = doc.createElement("a");
                        a.href = dataUrl;
                        a.download = fname;
                        a.rel = "noopener";
                        a.style.display = "none";
                        doc.body.appendChild(a);
                        a.click();
                        setTimeout(function () {{
                            try {{ a.remove(); }} catch (e) {{}}
                        }}, 200);
                        return true;
                    }} catch (e) {{
                        return false;
                    }}
                }}
                if (window.parent && window.parent !== window && tryDownload(window.parent.document)) return;
                if (window.top && window.top !== window && tryDownload(window.top.document)) return;
                tryDownload(document);
            }})();
        </script>
        """,
        height=0,
        width=0,
    )


def build_structured_protocol_txt(fields: dict, consultation_date: str) -> str:
    parts = [
        "Протокол консультации",
        "",
        f"Дата: {consultation_date}",
        "",
        "Жалобы",
        fields.get("complaints", "") or "—",
        "",
        "Анамнез",
        fields.get("anamnesis", "") or "—",
        "",
        "Заключение",
        fields.get("conclusion", "") or "—",
        "",
        "Рекомендации",
        fields.get("recommendations", "") or "—",
        "",
    ]
    return "\n".join(parts)


# ============ MAIN UI ============

st.title("🏥 ClinVoice")
st.markdown("**Распознавание речи для медицинских консультаций**")

hub_model_id = resolve_hub_model_id()
model_size = "small"

st.header("Запись консультации")
st.caption(
    "Чтобы сбросить поля приложения и буфер записи, используйте Clear session state "
    "(или кнопку «Сбросить запись…» на странице)."
)
if "original_transcription" not in st.session_state:
    st.session_state.original_transcription = None
if "protocol_editor_text" not in st.session_state:
    st.session_state.protocol_editor_text = ""


def _sync_live_transcript_from_whisper() -> None:
    ensure_webrtc_shared_initialized()
    if st.session_state.get("live_transcript_pause_auto_sync"):
        return
    lk = st.session_state.webrtc_shared.get("lock")
    if not lk:
        return
    with lk:
        live = st.session_state.webrtc_shared.get("live_whisper_text") or ""
    current = (st.session_state.get("live_transcript_editor") or "").strip()
    # После успешного финала буфер и черновик очищаются — не затирать поле пустым черновиком.
    if not live.strip() and current:
        return
    st.session_state.live_transcript_editor = live


_live_interval = resolve_live_whisper_interval_sec()
_asr_transcriber = get_cached_asr_transcriber(model_size, hub_model_id)

webrtc_streamer(
    key=CLINVOICE_WEBRTC_STREAMER_KEY,
    mode=WebRtcMode.SENDONLY,
    rtc_configuration=resolve_webrtc_rtc_configuration(),
    media_stream_constraints={"audio": True, "video": False},
    audio_processor_factory=build_webrtc_processor_factory(_asr_transcriber, _live_interval),
    async_processing=True,
    sendback_audio=False,
    translations=WEBRTC_UI_RU,
)

if st.session_state.pop("_pending_webrtc_full_reset", False):
    ensure_webrtc_shared_initialized()
    _rlk = st.session_state.webrtc_shared.get("lock")
    if _rlk:
        with _rlk:
            st.session_state.webrtc_shared["pcm_accum"] = bytearray()
            st.session_state.webrtc_shared["live_whisper_text"] = ""
            st.session_state.webrtc_shared["live_whisper_error"] = None
            st.session_state.webrtc_shared["live_draft_pcm_committed"] = 0
    st.session_state.live_transcript_editor = ""
    st.session_state.live_transcript_pause_auto_sync = False

if st.session_state.pop("_pending_apply_live_whisper", False):
    ensure_webrtc_shared_initialized()
    _alk = st.session_state.webrtc_shared.get("lock")
    _auto = ""
    if _alk:
        with _alk:
            _auto = st.session_state.webrtc_shared.get("live_whisper_text") or ""
    st.session_state.live_transcript_editor = _auto
    st.session_state.live_transcript_pause_auto_sync = False


@st.fragment(run_every=timedelta(milliseconds=400))
def _live_transcript_ui_fragment():
    """
    Один фрагмент с run_every: подтягивает live_whisper_text в поле и обновляет статус.
    Без on_change у text_area — иначе Streamlit может помечать поле как изменённое при
    программной подстановке и блокировать авто-синхронизацию.
    """
    _sync_live_transcript_from_whisper()
    st.checkbox(
        "Остановить авто-подстановку черновика (для ручного редактирования без перезаписи)",
        key="live_transcript_pause_auto_sync",
        help="Пока включено, текст в поле не обновляется из Whisper; используйте «Подставить последний авто-текст».",
    )
    st.text_area(
        "Транскрипт (можно править во время записи)",
        height=160,
        key="live_transcript_editor",
        help="Автоматически повторяет накопленный черновик Whisper, пока не включена остановка выше.",
    )
    if st.button("Подставить последний авто-текст", key="apply_live_whisper_to_editor"):
        st.session_state._pending_apply_live_whisper = True
        st.rerun()

    sh = st.session_state.webrtc_shared
    lk = sh.get("lock")
    sec = 0.0
    err: Optional[str] = None
    if lk:
        with lk:
            sec = len(sh.get("pcm_accum") or b"") / 32000.0
            err = sh.get("live_whisper_error")
    _iv = resolve_live_whisper_interval_sec()
    st.caption(
        f"Накоплено под Whisper: **~{sec:.1f}** с. Инкрементальный черновик — тик примерно каждые **{_iv:g}** с."
    )
    if err:
        st.error(err)


_live_transcript_ui_fragment()

if st.button("Сбросить запись и черновик", key="webrtc_reset_buffer"):
    st.session_state._pending_webrtc_full_reset = True
    st.rerun()

_max_seg_b = int(resolve_draft_tail_max_seconds() * 32000)
_overlap_b = int(resolve_draft_tail_overlap_sec() * 32000)

with st.expander("Дополнительно: полный Whisper по накопленному буферу", expanded=False):
    st.caption(
        "Один медленный прогон по всему PCM (beam 5). Результат для справки; основная кнопка ниже "
        "строит протокол по live-транскрипту и короткому догону хвоста."
    )
    if st.button("Запустить полный Whisper по буферу", key="full_buffer_whisper_once"):
        ensure_webrtc_shared_initialized()
        _sh = st.session_state.webrtc_shared
        _lk = _sh.get("lock")
        _al = _sh.get("asr_lock")
        _pcm = b""
        if _lk:
            with _lk:
                _pcm = bytes(_sh.get("pcm_accum") or b"")
        if len(_pcm) < 3200:
            st.warning("В буфере слишком мало аудио (меньше ~0,1 с).")
        else:
            _path = None
            try:
                fd, _path = tempfile.mkstemp(suffix=".wav")
                os.close(fd)
                with open(_path, "wb") as wf:
                    wf.write(pcm_mono_s16le_to_wav_bytes(_pcm))
                with st.spinner("Полная транскрибация…"):
                    if _al:
                        with _al:
                            _full_txt = transcribe_wav_in_chunks(
                                _asr_transcriber, _path, language="ru", draft=False
                            )
                    else:
                        _full_txt = transcribe_wav_in_chunks(
                            _asr_transcriber, _path, language="ru", draft=False
                        )
                st.code(_full_txt or "—", language=None)
            except Exception as e:
                st.error(str(e))
            finally:
                if _path and os.path.isfile(_path):
                    try:
                        os.remove(_path)
                    except OSError:
                        pass

if st.button("Заполнить протокол по транскрипту", type="primary"):
    ensure_webrtc_shared_initialized()
    if not yandex_llm_configured():
        st.error(
            "Не заданы параметры для заполнения протокола: укажите "
            "**YANDEX_FOLDER_ID** и (**YANDEX_CLOUD_API_KEY** или **YANDEX_IAM_TOKEN**) "
            "в переменных окружения или в секретах приложения."
        )
        st.stop()

    sh = st.session_state.webrtc_shared
    lk = sh.get("lock")
    asr_lock = sh.get("asr_lock") or threading.Lock()

    with st.spinner("Дораспознавание хвоста записи…"):
        flush_err = flush_incremental_pcm_tail(
            sh, _asr_transcriber, asr_lock, _max_seg_b, _overlap_b
        )
    if flush_err:
        st.error(flush_err)
        st.stop()

    if not st.session_state.get("live_transcript_pause_auto_sync") and lk:
        with lk:
            _live_sync = (st.session_state.webrtc_shared.get("live_whisper_text") or "").strip()
        st.session_state.live_transcript_editor = _live_sync

    editor_raw = (st.session_state.get("live_transcript_editor") or "").strip()
    live_raw = ""
    if lk:
        with lk:
            live_raw = (sh.get("live_whisper_text") or "").strip()
    transcription = editor_raw if editor_raw else live_raw
    if not transcription:
        st.error(
            "Нет текста для протокола: сначала запишите консультацию (или вставьте текст в поле транскрипта)."
        )
        st.stop()

    st.session_state.original_transcription = transcription
    st.session_state.doctor_transcript_editor = transcription
    st.session_state.protocol_consultation_date = format_consultation_date_gmt3()

    try:
        with st.spinner("Заполнение протокола..."):
            protocol = fill_protocol_from_transcript(transcription)
        st.session_state.protocol_editor_text = format_protocol_editor_text(
            st.session_state.protocol_consultation_date,
            protocol,
        )
        st.success(
            "Готово. Проверьте и при необходимости отредактируйте протокол ниже. "
            "Текст протокола должен сохраниться автоматически как **протокол.txt**; "
            "если загрузка не началась, скачайте файл кнопкой в блоке протокола."
        )
        _auto_txt = build_structured_protocol_txt(protocol, st.session_state.protocol_consultation_date)
        trigger_browser_text_download("протокол.txt", _auto_txt)
        if lk:
            with lk:
                st.session_state.webrtc_shared["pcm_accum"] = bytearray()
                st.session_state.webrtc_shared["live_whisper_text"] = ""
                st.session_state.webrtc_shared["live_whisper_error"] = None
                st.session_state.webrtc_shared["live_draft_pcm_committed"] = 0
        st.session_state.live_transcript_pause_auto_sync = False
    except Exception as e:
        st.error(f"Ошибка заполнения протокола: {e}")
        st.session_state.protocol_editor_text = format_protocol_editor_text(
            st.session_state.protocol_consultation_date,
            {k: "" for k in PROTOCOL_FIELD_KEYS},
        )

if st.session_state.original_transcription:
    if "doctor_transcript_editor" not in st.session_state:
        st.session_state.doctor_transcript_editor = st.session_state.original_transcription

    st.header("Протокол консультации")

    with st.expander("Транскрипт (справочно, можно исправить ошибки распознавания)"):
        st.text_area(
            "Транскрипт",
            height=180,
            key="doctor_transcript_editor",
            label_visibility="collapsed",
        )

    st.text_area(
        "Протокол (редактирование)",
        height=320,
        key="protocol_editor_text",
        help="Формат: строки «Дата:», «Жалобы:», «Анамнез:», «Заключение:», «Рекомендации:» с двоеточием.",
    )

    _parsed_date, fields = parse_protocol_editor_text(
        st.session_state.get("protocol_editor_text", "")
    )
    consultation_date = (
        _parsed_date.strip()
        or st.session_state.get("protocol_consultation_date")
        or format_consultation_date_gmt3()
    )

    st.subheader("Скачать протокол")
    doc = create_structured_protocol_docx(fields, consultation_date)
    doc_buf = io.BytesIO()
    doc.save(doc_buf)
    doc_bytes = doc_buf.getvalue()
    txt_body = build_structured_protocol_txt(fields, consultation_date)

    col_docx, col_txt = st.columns(2)
    col_docx.download_button(
        "Скачать протокол (.docx)",
        doc_bytes,
        "протокол.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    col_txt.download_button(
        "Скачать протокол (.txt)",
        txt_body,
        "протокол.txt",
        "text/plain",
    )

# Footer
st.markdown("---")
st.markdown("*ClinVoice v0.4 — распознавание речи для консультаций*")
